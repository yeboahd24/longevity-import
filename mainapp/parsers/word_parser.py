import re
import fitz
import tabula
from langdetect import detect, lang_detect_exception
from collections import Counter
from translate import Translator
from docx import Document

from mainapp.parsers.constants import KEYWORD_UNITS, KEYWORD_COMMENT


class WordFileParser:
    def __init__(self, file_path):
        self.doc = Document(file_path)
        self.data_dict = dict()
        self.language = self.detect_language()

    def detect_language(self):
        """Determine the document language."""

        lang_list = list()
        for table in self.doc.tables:
            for i, row in enumerate(table.rows):
                tuple_text = tuple(cell.text for cell in row.cells)
                for text in tuple_text:
                    text = re.sub(r'[\.,\/#!$%\^&\*;:{}=\-–—_`~()]', '', text)
                    try:
                        lang_list.append(detect(text))
                    except lang_detect_exception.LangDetectException:
                        pass

        # convert the list with languages into a Counter and get
        # the most common language
        counter_lang = Counter(lang_list)
        language = counter_lang.most_common(1)[0][0]
        return language

    def get_analyzes_table(self, units_on_lang):
        """Get the analysis table."""

        # from the list of tables obtained, determine
        # which is the table with keyword analyses
        for table in self.doc.tables:
            headers = table.rows[0]
            text_tuple = tuple(cell.text for cell in headers.cells)
            for i in units_on_lang:
                res = [head for head in text_tuple if head.lower() == i.lower()]
                if res:
                    print()
                    return table

    def _get_index_col(self, current_table, values):
        """Get the index of the column."""

        # find the index of the column by keyword
        headers = current_table.rows[0]
        text_tuple = tuple(cell.text for cell in headers.cells)
        for i in values:
            col_name = [head for head in text_tuple if head.lower() == i.lower()][0]
            col_index = text_tuple.index(col_name)
            return col_index

    def _get_index_reference_col(self, current_table):
        """Get the index of the reference values column."""

        # we find the index of the column of reference
        # values according to the regular expression pattern
        table_data = current_table.rows[1:]
        res_list = list()
        for row in table_data:
            for index, j in enumerate(row.cells):
                if re.fullmatch(r'\d+\.\d+\s[\-–—]\s\d+\.\d+', j.text):
                    res_list.append(index)
        counter_lang = Counter(res_list)
        index_col = counter_lang.most_common(1)[0][0]
        return index_col

    def _get_analyzes(self, current_table, comment_col_index, units_col_index, reference_col_index):
        """Get analysis data."""

        analyzes_list = list()
        translator = Translator(from_lang=self.language, to_lang='en')
        analyzes_row_index = 1
        while True:
            item_dict = dict()

            table_row = current_table.rows[analyzes_row_index]

            # read the name of the analysis from the first row of the
            # analysis table, if empty, go to the next row
            param_name = table_row.cells[0].text
            if param_name:
                param_name = translator.translate(param_name)
                param_name = param_name.lower()
                item_dict['parameter'] = param_name
            else:
                analyzes_row_index += 1
                continue

            # read the result of the analysis
            param_value = table_row.cells[analyzes_row_index].text
            if param_name:
                item_dict['value'] = param_value
            else:
                item_dict['value'] = None

            # read the comment to the analysis
            if comment_col_index:
                param_comment = table_row.cells[comment_col_index].text
                if param_comment:
                    param_comment = translator.translate(param_comment)
                    item_dict['comment'] = param_comment
                else:
                    item_dict['comment'] = None
            else:
                item_dict['comment'] = None

            # read the units of measurement of the analysis result
            units = table_row.cells[units_col_index].text
            if units:
                units = translator.translate(units)
                item_dict['units'] = units
            else:
                item_dict['units'] = None

            # read and parse the field of reference values
            reference_value = table_row.cells[reference_col_index].text
            reference_values = self._parse_ref_value(reference_value)
            item_dict['reference_values'] = reference_values

            # adding a dictionary with the received data to the list of analyses
            analyzes_list.append(item_dict)
            analyzes_row_index += 1

            # if we get an error when accessing the units of measurement
            # field of the next line, we exit the loop
            try:
                current_table.rows[analyzes_row_index]
            except IndexError:
                break
        return analyzes_list

    def _parse_ref_value(self, value):
        """Parse the string reference values."""

        min_value = None
        max_value = None
        if value:
            res = re.search('[\-–—]', value)
            if res:
                value_list = re.split('[\-–—]', value)
                min_value = value_list[0].strip()
                max_value = value_list[1].strip()
            else:
                if value.find('<') != -1:
                    max_value = re.findall('\d.+', value)
                elif value.find('>') != -1:
                    min_value = re.findall('\d+\.?\d+', value)
        return min_value, max_value

    def run(self):
        """Run file parsing."""

        self.data_dict['language'] = self.language

        # creating a translator object to translate
        # data into the document language
        translator = Translator(to_lang=self.language)
        units_on_lang = translator.translate(KEYWORD_UNITS).split(' ')
        comment_on_lang = translator.translate(KEYWORD_COMMENT).split(' ')

        current_table = self.get_analyzes_table(units_on_lang)

        comments_col_index = self._get_index_col(current_table, comment_on_lang)
        units_col_index = self._get_index_col(current_table, units_on_lang)
        reference_col_index = self._get_index_reference_col(current_table)

        analyzes_list = self._get_analyzes(
            current_table,
            comments_col_index,
            units_col_index,
            reference_col_index)

        self.data_dict['analyzes'] = analyzes_list

    def get_data(self):
        """Get the parsing result."""

        return self.data_dict
